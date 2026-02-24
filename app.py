"""
app.py - 글로벌 뉴스 클리핑 Streamlit 앱
실행: streamlit run app.py
"""

import streamlit as st
from datetime import datetime
import html as html_lib
import subprocess, sys

from crawler import run_pipeline, SOURCES, KEYWORD_TRANSLATIONS

# ─── 페이지 설정 ──────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="글로벌 뉴스 클리핑",
    page_icon="📰",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── CSS ─────────────────────────────────────────────────────────────────────

st.markdown("""
<style>
html, body, [class*="css"] { font-family: 'Noto Sans KR', sans-serif; }

.main-header {
    background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
    color: white; padding: 2rem 2.5rem; border-radius: 12px; margin-bottom: 1.5rem;
}
.main-header h1 { margin: 0; font-size: 1.8rem; font-weight: 700; }
.main-header p  { margin: 0.4rem 0 0; opacity: 0.75; font-size: 0.9rem; }

.article-card {
    background: #ffffff; border: 1px solid #e8e8e8;
    border-left: 4px solid #0f3460; border-radius: 8px;
    padding: 0.85rem 1rem; margin-bottom: 0.7rem;
    transition: box-shadow 0.2s;
}
.article-card:hover { box-shadow: 0 3px 12px rgba(0,0,0,0.1); }

.article-title-ko { font-size: 0.95rem; font-weight: 600; color: #1a1a2e; margin-bottom: 0.2rem; }
.article-title-ko a { text-decoration: none; color: inherit; }
.article-title-ko a:hover { color: #0f3460; text-decoration: underline; }
.article-title-orig { font-size: 0.78rem; color: #888; margin-bottom: 0.35rem; }
.article-meta { display: flex; gap: 0.5rem; flex-wrap: wrap; align-items: center; }
.badge      { background: #f0f4ff; color: #0f3460; border-radius: 4px; padding: 1px 8px; font-size: 0.72rem; font-weight: 500; }
.badge-date { background: #f5f5f5; color: #666;    border-radius: 4px; padding: 1px 8px; font-size: 0.72rem; }

.stat-box   { background: #f8f9ff; border: 1px solid #dce3ff; border-radius: 8px; padding: 0.6rem 1rem; text-align: center; }
.stat-num   { font-size: 1.6rem; font-weight: 700; color: #0f3460; }
.stat-label { font-size: 0.75rem; color: #666; }

.empty-state { text-align: center; color: #aaa; padding: 2rem; font-size: 0.9rem; }
</style>
""", unsafe_allow_html=True)

# ─── 헬퍼 ────────────────────────────────────────────────────────────────────

def safe(text: str) -> str:
    return html_lib.escape(str(text or ""))

def render_card(article: dict):
    title_ko   = safe(article.get("title_ko") or article.get("title", ""))
    title_orig = safe(article.get("title", ""))
    url        = html_lib.escape(article.get("url", "#"))
    source     = safe(article.get("source", ""))
    flag       = article.get("flag", "")
    date       = article.get("date", "")
    date_str   = date[:10] if date else ""
    date_badge = f'<span class="badge-date">📅 {date_str}</span>' if date_str else ""

    st.markdown(f"""
    <div class="article-card">
        <div class="article-title-ko"><a href="{url}" target="_blank">{title_ko}</a></div>
        <div class="article-title-orig">{title_orig}</div>
        <div class="article-meta">
            <span class="badge">{flag} {source}</span>
            {date_badge}
        </div>
    </div>
    """, unsafe_allow_html=True)


# ─── Word 문서 생성 ───────────────────────────────────────────────────────────

def build_docx(result: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    import io

    meta = result["meta"]
    doc  = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    title_p = doc.add_heading("글로벌 뉴스 클리핑 리포트", level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run(
        f"키워드: {meta['keyword']}  |  생성일시: {meta['generated_at']}  |  수집 기간: 최근 {meta['days']}일"
    ).font.size = Pt(9)
    doc.add_paragraph()

    def add_section(heading: str, articles: list):
        doc.add_heading(heading, level=1)
        if not articles:
            doc.add_paragraph("수집된 기사가 없습니다.").italic = True
            doc.add_paragraph()
            return
        for a in articles:
            title_ko   = a.get("title_ko") or a.get("title", "")
            title_orig = a.get("title", "")
            url        = a.get("url", "")
            source     = a.get("source", "")
            flag       = a.get("flag", "")
            date       = (a.get("date", "") or "")[:10]

            p   = doc.add_paragraph(style="List Bullet")
            run = p.add_run(title_ko)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)

            if title_orig and title_orig != title_ko:
                orig_p = doc.add_paragraph()
                orig_p.paragraph_format.left_indent = Inches(0.3)
                r = orig_p.add_run(f"원문: {title_orig}")
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.left_indent = Inches(0.3)
            meta_r = meta_p.add_run(f"{flag} {source}" + (f"  |  {date}" if date else "") + f"\n{url}")
            meta_r.font.size = Pt(8)
            meta_r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            doc.add_paragraph()

    add_section("🇯🇵 일본 산업 이슈", result["japan"])
    add_section("🇨🇳 중국 산업 이슈", result["china"])
    add_section("🇹🇼 대만 산업 이슈", result["taiwan"])

    doc.add_paragraph()
    footer_p = doc.add_paragraph("* 본 리포트는 자동 수집·Google 번역된 내용으로, 원문 확인을 권장합니다.")
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── 사이드바 ─────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown("## ⚙️ 검색 설정")

    preset_keywords = list(KEYWORD_TRANSLATIONS.keys()) + ["직접 입력"]
    selected_preset = st.selectbox("키워드 프리셋", preset_keywords, index=1)
    if selected_preset == "직접 입력":
        keyword = st.text_input("키워드 직접 입력", placeholder="예: 패션 브랜드")
    else:
        keyword = selected_preset

    st.divider()
    days = st.slider("수집 기간 (일)", min_value=1, max_value=30, value=7)
    st.divider()

    st.markdown("**수집 매체 선택**")
    use_japan  = st.checkbox("🇯🇵 일본",  value=True)
    use_china  = st.checkbox("🇨🇳 중국",  value=True)
    use_taiwan = st.checkbox("🇹🇼 대만",  value=True)
    st.divider()

    run_btn = st.button("🔍 뉴스 수집 시작", type="primary", use_container_width=True)

    st.divider()
    st.markdown("""
    <div style="font-size:0.75rem;color:#999;line-height:1.6;">
    ℹ️ Google 번역(비공식 무료 API) 사용<br>
    ⚠️ 일부 매체는 크롤링 제한으로 수집 불가<br>
    📅 날짜 미확인 기사는 자동 제외됩니다
    </div>
    """, unsafe_allow_html=True)

# ─── 헤더 ────────────────────────────────────────────────────────────────────

st.markdown("""
<div class="main-header">
    <h1>📰 글로벌 뉴스 클리핑</h1>
    <p>일본 · 중국 · 대만 패션/리테일 매체 자동 수집 & 한국어 번역</p>
</div>
""", unsafe_allow_html=True)

# ─── 실행 ────────────────────────────────────────────────────────────────────

if "result" not in st.session_state:
    st.session_state.result = None

if run_btn:
    if not keyword or not keyword.strip():
        st.warning("키워드를 입력해주세요.")
        st.stop()

    active = []
    if use_japan:  active.append("japan")
    if use_china:  active.append("china")
    if use_taiwan: active.append("taiwan")
    if not active:
        st.warning("최소 한 개의 매체 국가를 선택해주세요.")
        st.stop()

    status_ph = st.empty()
    prog_bar  = st.progress(0)
    prog_text = st.empty()

    try:
        result = run_pipeline(
            keyword_ko=keyword.strip(),
            days=days,
            active_categories=active,
            on_status=lambda msg: status_ph.info(f"⏳ {msg}"),
            on_progress=lambda v, t: (prog_bar.progress(v), prog_text.caption(t)),
        )
        st.session_state.result = result
        status_ph.empty(); prog_bar.empty(); prog_text.empty()
        st.success(f"✅ 수집 완료! — {result['meta']['generated_at']}")
    except Exception as e:
        status_ph.error(f"오류 발생: {e}")
        prog_bar.empty()

# ─── 결과 표시 ────────────────────────────────────────────────────────────────

if st.session_state.result:
    result = st.session_state.result
    meta   = result["meta"]

    # 통계 (총 수집 + 국가별 3개 = 컬럼 4개)
    counts = {k: len(result[k]) for k in ["japan", "china", "taiwan"]}
    total  = sum(counts.values())
    c1, c2, c3, c4 = st.columns(4)
    for col, (label, key) in zip(
        [c1, c2, c3, c4],
        [("총 수집", None), ("🇯🇵 일본", "japan"), ("🇨🇳 중국", "china"), ("🇹🇼 대만", "taiwan")]
    ):
        n = total if key is None else counts[key]
        col.markdown(
            f'<div class="stat-box"><div class="stat-num">{n}</div>'
            f'<div class="stat-label">{label}</div></div>',
            unsafe_allow_html=True
        )

    st.markdown("<br>", unsafe_allow_html=True)

    # Word 다운로드
    date_str = datetime.now().strftime("%Y%m%d_%H%M")
    try:
        docx_bytes = build_docx(result)
        st.download_button(
            label="⬇️ 리포트 다운로드 (.docx)",
            data=docx_bytes,
            file_name=f"news_clipping_{date_str}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        st.warning(f"Word 문서 생성 실패: {e}. requirements.txt에 python-docx가 있는지 확인해주세요.")

    st.divider()

    # 탭 (국가별 3개)
    t1, t2, t3 = st.tabs([
        f"🇯🇵 일본 ({counts['japan']})",
        f"🇨🇳 중국 ({counts['china']})",
        f"🇹🇼 대만 ({counts['taiwan']})",
    ])

    with t1:
        if result["japan"]:
            for a in result["japan"]: render_card(a)
        else:
            st.markdown('<div class="empty-state">🔍 수집된 일본 기사가 없습니다.</div>', unsafe_allow_html=True)

    with t2:
        if result["china"]:
            for a in result["china"]: render_card(a)
        else:
            st.markdown('<div class="empty-state">🔍 수집된 중국 기사가 없습니다.</div>', unsafe_allow_html=True)

    with t3:
        if result["taiwan"]:
            for a in result["taiwan"]: render_card(a)
        else:
            st.markdown('<div class="empty-state">🔍 수집된 대만 기사가 없습니다.</div>', unsafe_allow_html=True)

else:
    st.markdown("""
    <div style="text-align:center;padding:4rem 2rem;color:#aaa;">
        <div style="font-size:3rem;margin-bottom:1rem;">🌏</div>
        <div style="font-size:1.1rem;font-weight:600;color:#555;margin-bottom:0.5rem;">
            좌측 사이드바에서 키워드와 옵션을 설정한 뒤<br>
            <strong>뉴스 수집 시작</strong> 버튼을 눌러주세요.
        </div>
        <div style="font-size:0.85rem;margin-top:1rem;">
            일본 · 중국 · 대만 매체에서 기사를 수집하고<br>
            Google 번역으로 한국어 제목을 자동 생성합니다.
        </div>
    </div>
    """, unsafe_allow_html=True)
